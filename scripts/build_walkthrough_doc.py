"""Generate PROJECT_WALKTHROUGH.docx — a plain-language tour of the project.

Reads models/metrics.json (if present) so the numbers in the doc reflect the
most recent training run.

Usage:
    python scripts/build_walkthrough_doc.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parents[1]
METRICS_PATH = ROOT / "models" / "metrics.json"
OUT_PATH = ROOT / "PROJECT_WALKTHROUGH.docx"


def _load_metrics() -> dict | None:
    if METRICS_PATH.exists():
        try:
            return json.loads(METRICS_PATH.read_text())
        except Exception:
            return None
    return None


def _style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def _add_para(doc, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _add_bullets(doc, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _add_code(doc, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _metrics_table(doc, metrics: dict | None) -> None:
    if metrics is None:
        _add_para(
            doc,
            "(No models/metrics.json found yet — run `python -m diabetes_mlops.train` to populate.)",
        )
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    rows = [
        ("Selected model", str(metrics.get("selected", "—"))),
        ("ROC-AUC (test)", f"{metrics.get('roc_auc', 0):.3f}"),
        ("Accuracy (test)", f"{metrics.get('accuracy', 0):.3f}"),
        ("Recall, positive class (test)", f"{metrics.get('recall', 0):.3f}"),
        ("F1 (test)", f"{metrics.get('f1', 0):.3f}"),
        ("Decision threshold", f"{metrics.get('threshold', 0):.3f}"),
        ("CV out-of-fold ROC-AUC", f"{metrics.get('cv_oof_roc_auc', 0):.3f}"),
    ]
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value


def build() -> Path:
    metrics = _load_metrics()
    doc = Document()
    _style(doc)

    # ─── Title ─────────────────────────────────────────────────────────
    title = doc.add_heading("Diabetes Risk Predictor", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    r = sub.add_run("End-to-End MLOps Walkthrough — from modelling to deployment")
    r.italic = True
    r.font.size = Pt(13)

    doc.add_paragraph(
        "Author: Tadaishe Maumbe   |   Stack: Python · scikit-learn · MLflow · "
        "FastAPI · Streamlit · Docker · GitHub Actions"
    )

    # ─── 1. The story ────────────────────────────────────────────────
    _add_heading(doc, "1. The Story in One Paragraph", 1)
    doc.add_paragraph(
        "This project takes a clinical screening question — 'which patients walking into a "
        "primary-care clinic should be prioritised for a confirmatory diabetes test?' — and "
        "turns it into a working web app that any clinician can use. It is deliberately a "
        "full-stack MLOps project, not just a notebook: data ingestion, cleaning, training, "
        "experiment tracking, two deployment surfaces (a Streamlit UI and a FastAPI REST "
        "service), containerisation with Docker, and continuous integration with GitHub "
        "Actions. The point is to show that a model can move out of a Jupyter notebook and "
        "into something usable."
    )

    # ─── 2. The business problem ─────────────────────────────────────
    _add_heading(doc, "2. The Business / Clinical Problem", 1)
    doc.add_paragraph(
        "Type-2 diabetes affects more than 530 million adults worldwide, and roughly 45% of "
        "cases go undiagnosed until complications appear. In under-resourced clinical "
        "settings the gap is even wider because of limited access to specialist testing. "
        "The model in this project is a screening triage tool: given eight routine "
        "clinical measurements, it estimates the probability that a patient has diabetes, "
        "so a clinician can decide whether to order a confirmatory HbA1c."
    )
    _add_para(doc, "Key framing decisions:", bold=True)
    _add_bullets(
        doc,
        [
            "This is screening, not diagnosis. The downstream action is 'order another test', not 'prescribe insulin'.",
            "False negatives (missed diabetic patients) cost more than false positives (an unnecessary follow-up). The decision threshold reflects that.",
            "The model must be explainable enough that a clinician can interrogate any single prediction — otherwise it will not be trusted.",
        ],
    )

    # ─── 3. The data ─────────────────────────────────────────────────
    _add_heading(doc, "3. The Data", 1)
    doc.add_paragraph(
        "Pima Indians Diabetes Database (US National Institute of Diabetes and Digestive and "
        "Kidney Diseases). 768 female patients aged 21+, 8 numeric features, binary outcome. "
        "Public mirror; auto-downloaded by `python -m diabetes_mlops.data`."
    )
    _add_para(doc, "Domain-aware data quality note:", bold=True)
    doc.add_paragraph(
        "Five columns (Glucose, BloodPressure, SkinThickness, Insulin, BMI) encode missing "
        "values as a literal zero — physiologically impossible. A domain-naive pipeline "
        "would carry those zeros into training and quietly learn that 'zero means high risk'. "
        "`data.clean()` replaces them with NaN; the sklearn pipeline's SimpleImputer then "
        "fills them in inside cross-validation folds, so no test-fold information leaks."
    )

    # ─── 4. Step-by-step pipeline ────────────────────────────────────
    _add_heading(doc, "4. What Actually Happens, Step by Step", 1)

    _add_heading(doc, "4.1  Ingest", 2)
    doc.add_paragraph(
        "`src/diabetes_mlops/data.py` downloads the CSV (only if not already on disk), "
        "attaches the canonical column headers, and writes data/raw/diabetes.csv. "
        "`build_processed()` then runs `clean()` and writes data/processed/diabetes_clean.csv. "
        "Raw and processed files are gitignored — the code is the source of truth, not the "
        "binary."
    )

    _add_heading(doc, "4.2  Feature pipeline", 2)
    doc.add_paragraph(
        "`features.make_preprocessor()` returns a scikit-learn Pipeline of two steps: median "
        "imputation followed by standard scaling. Wrapping these in a Pipeline (rather than "
        "doing them as one-off pandas transforms) means the exact same code path runs at "
        "training time and at inference time — no train/serve skew."
    )

    _add_heading(doc, "4.3  Train and evaluate", 2)
    doc.add_paragraph(
        "`train.train()` does the following for each candidate model (Logistic Regression, "
        "Random Forest, HistGradientBoosting):"
    )
    _add_bullets(
        doc,
        [
            "Stratified 5-fold cross-validation on the training partition, collecting out-of-fold predicted probabilities.",
            "Threshold tuning: find the largest decision threshold that still achieves recall ≥ 0.85 on those OOF predictions (the clinician-set policy parameter).",
            "Refit the full pipeline on the entire training partition.",
            "Score on the held-out 20% test set — ROC-AUC, accuracy, recall, F1, confusion matrix.",
            "Log everything to MLflow: parameters, metrics, the fitted pipeline, the run name.",
        ],
    )
    doc.add_paragraph(
        "After all candidates are evaluated, the model with the highest ROC-AUC is "
        "promoted: it is saved (together with its threshold, feature order, and metrics) to "
        "models/diabetes_pipeline.joblib, and a flat summary is written to "
        "models/metrics.json. Both deployment surfaces consume the same joblib file."
    )

    _add_heading(doc, "4.4  Experiment tracking with MLflow", 2)
    doc.add_paragraph(
        "Every training run shows up in MLflow. Launch the UI with `mlflow ui` from the "
        "project root and open http://127.0.0.1:5000. You can:"
    )
    _add_bullets(
        doc,
        [
            "Compare any two runs side-by-side (e.g. baseline LR vs the gradient-boosted final).",
            "Filter by metric — show only runs with ROC-AUC > 0.85.",
            "Open the artifact tab to download the exact pipeline that produced a given metric.",
            "Tag a run with the git commit SHA — links the metrics back to the code that made them.",
        ],
    )

    # ─── 5. Deployment ───────────────────────────────────────────────
    _add_heading(doc, "5. Deployment Surfaces", 1)

    _add_heading(doc, "5.1  Streamlit web app", 2)
    doc.add_paragraph(
        "`app/streamlit_app.py` exposes the model as a clinician-facing UI. Sliders for "
        "each input feature, a progress-bar risk gauge, and a coloured verdict — green if "
        "below the screening threshold, red if the patient should be flagged for follow-up. "
        "The sidebar shows the model algorithm, threshold, and headline metrics so the user "
        "always knows what they are looking at."
    )
    _add_code(doc, "streamlit run app/streamlit_app.py")

    _add_heading(doc, "5.2  FastAPI service", 2)
    doc.add_paragraph(
        "`api/main.py` exposes the same model as a REST API with two endpoints:"
    )
    _add_bullets(
        doc,
        [
            "GET /health → returns the loaded model name; used by container orchestrators for liveness checks.",
            "POST /predict → takes a JSON body of patient measurements, returns probability, label, threshold, and model name.",
        ],
    )
    doc.add_paragraph(
        "Pydantic models on the request side give automatic validation (e.g. BMI must be "
        "between 0 and 80) and free OpenAPI documentation at /docs."
    )
    _add_code(doc, "uvicorn api.main:app --reload")

    _add_heading(doc, "5.3  Docker", 2)
    doc.add_paragraph(
        "The Dockerfile installs only what is needed for serving (no notebooks, no test "
        "data) and starts uvicorn on port 8000. One image, deployable to any container "
        "platform — Fly.io, Render, AWS App Runner, Azure Container Apps."
    )
    _add_code(doc, "docker build -t diabetes-risk .\ndocker run -p 8000:8000 diabetes-risk")

    # ─── 6. CI ─────────────────────────────────────────────────────
    _add_heading(doc, "6. Continuous Integration", 1)
    doc.add_paragraph(
        "`.github/workflows/ci.yml` runs on every push and pull request. It installs the "
        "dependencies, lints with ruff, and runs the pytest suite — including a smoke "
        "training run that proves the entire pipeline still works end-to-end. If the smoke "
        "train succeeds, the resulting model artifact is uploaded so it can be inspected "
        "from the Actions tab."
    )

    # ─── 7. Results ─────────────────────────────────────────────────
    _add_heading(doc, "7. Results from the Local Run", 1)
    _metrics_table(doc, metrics)
    if metrics is not None:
        doc.add_paragraph(
            "These are the headline numbers from the held-out 20% test set, with the "
            "decision threshold chosen on cross-validated out-of-fold predictions. The "
            "selected model and threshold are stored inside the joblib artifact and applied "
            "automatically by both the Streamlit app and the FastAPI service."
        )

    # ─── 8. What's deliberately not done ────────────────────────────
    _add_heading(doc, "8. What This Project Deliberately Is NOT", 1)
    _add_bullets(
        doc,
        [
            "Not a diagnostic device. It is a screening probability, not a clinical diagnosis. Any deployment in a real care setting would require regulatory approval and prospective validation.",
            "Not generalisable beyond its training population. The Pima dataset is female, aged 21+, of Pima heritage. Performance on other populations would need re-validation.",
            "Not yet calibrated. The next obvious step is CalibratedClassifierCV (isotonic) so the predicted probabilities can be read directly as clinical risk.",
            "Not monitored in production. A real deployment would add drift detection (e.g. Evidently or NannyML) and a feedback loop from confirmed HbA1c results.",
        ],
    )

    # ─── 9. How to reproduce ────────────────────────────────────────
    _add_heading(doc, "9. How to Reproduce", 1)
    _add_code(
        doc,
        "git clone https://github.com/<your-user>/diabetes-risk-mlops.git\n"
        "cd diabetes-risk-mlops\n"
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "pip install -r requirements.txt\n"
        "python -m diabetes_mlops.data\n"
        "python -m diabetes_mlops.train\n"
        "mlflow ui                          # http://127.0.0.1:5000\n"
        "streamlit run app/streamlit_app.py # http://localhost:8501\n"
        "uvicorn api.main:app --reload      # http://127.0.0.1:8000/docs",
    )

    # ─── 10. Repository layout ──────────────────────────────────────
    _add_heading(doc, "10. Repository Layout", 1)
    _add_code(
        doc,
        "diabetes-risk-mlops/\n"
        "├── README.md                business problem → methodology → results\n"
        "├── requirements.txt         pinned dependency floor\n"
        "├── Dockerfile               container image for the API\n"
        "├── pyproject.toml           build config + ruff + pytest config\n"
        "├── .github/workflows/ci.yml CI on every push/PR\n"
        "├── src/diabetes_mlops/      library: config, data, features, train, predict\n"
        "├── app/streamlit_app.py     clinician-facing web UI\n"
        "├── api/main.py              FastAPI REST service\n"
        "├── tests/                   pytest suite (data + end-to-end smoke)\n"
        "├── notebooks/01_eda.ipynb   exploration notebook (not source of truth)\n"
        "├── docs/                    methodology.md, model-card.md, results.md\n"
        "├── data/{raw,processed}/    auto-populated, gitignored\n"
        "├── models/                  joblib artifact + metrics.json, gitignored\n"
        "└── mlruns/                  MLflow experiment store, gitignored",
    )

    # ─── 11. Skills demonstrated ────────────────────────────────────
    _add_heading(doc, "11. Skills Demonstrated", 1)
    _add_bullets(
        doc,
        [
            "Pipeline creation — scikit-learn Pipeline keeps training and inference paths identical.",
            "Model deployment — same artifact powers a Streamlit UI and a FastAPI REST service.",
            "API development — Pydantic-validated request bodies, OpenAPI docs at /docs.",
            "Experiment tracking — MLflow logs every run (parameters, metrics, artifacts, threshold).",
            "Reproducibility — deterministic seed, CI smoke training run, locked requirements.",
            "Containerisation — slim Dockerfile, .dockerignore for fast builds.",
            "Domain knowledge — clinically motivated cleaning, threshold tuning, and explainability requirements.",
        ],
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out}")
